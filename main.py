from tabula import read_pdf
import PyPDF2
import docx
from docx.shared import Pt

pdf_1 = open("Sample.pdf",'rb')
pdfReader = PyPDF2.PdfFileReader(pdf_1)

pageObj = []

for i in range(0,pdfReader.numPages):
    pageObj = pdfReader.getPage(i)
    page_text = pageObj.extractText()

    doc = docx.Document()
    string = "PDF Page "+str(i+1)
    doc.add_heading(string, 0)
    doc.add_paragraph(page_text)
    doc_name = string+".docx"
    doc.save(doc_name)
    
pdf_1.close()