import docx
import os



def Single_Page(pdfReader):
    pageObj = []
    doc = docx.Document()
    cd = os.chdir('./PDF_To_Page/')
    print(os.getcwd())
    for i in range(0,pdfReader.numPages):
        pageObj = pdfReader.getPage(i)
        page_text = pageObj.extractText()
        string = "PDF Page "+str(i+1)
        doc.add_heading(string, 0)
        doc.add_paragraph(page_text)
        doc.add_page_break()


    string = "Word_Doc-Converted"
    doc_name = string+".docx"
    doc.save(doc_name)
    return "Success"