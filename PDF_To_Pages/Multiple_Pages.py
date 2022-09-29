import docx
import os




def Multiple_Pages(pdfReader):
    pageObj = []
    cd = os.chdir('./PDF_To_Pages/')
    print(os.getcwd())
    for i in range(0,pdfReader.numPages):
        pageObj = pdfReader.getPage(i)
        page_text = pageObj.extractText()

        doc = docx.Document()
        string = "PDF Page "+str(i+1)
        doc.add_heading(string, 0)
        doc.add_paragraph(page_text)
        doc_name = string+".docx"
        doc.save(doc_name)
    
    return "Success"